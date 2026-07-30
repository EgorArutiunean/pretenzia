from __future__ import annotations

import argparse
import re
import shutil
import sys
import uuid
from dataclasses import dataclass, replace
from datetime import date, datetime
from decimal import Decimal, ROUND_FLOOR, ROUND_HALF_UP
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill

from app.modules.claims.generate_claims_from_registry import (
    ClaimIssue,
    ClaimRow,
    format_money,
    normalize_header as normalize_registry_header,
    parse_money,
    read_registry,
)
from app.shared.file_utils import safe_filename
from app.shared.money_to_words import money_to_words
from app.shared.zip_utils import create_zip


REQUIRED_BASE_COLUMNS = [
    "Номер объекта",
    "Адрес",
    "Компания",
    "Реквизиты",
]

REQUIRED_JURISDICTION_COLUMNS = [
    "Номер объекта",
    "Подсудность",
]

COURT_ORDER_FONT_NAME = "Times New Roman"
COURT_ORDER_FONT_SIZE_PT = 12
LIMITATION_PERIOD_MONTHS = 36
MIN_COURT_ORDER_AMOUNT = Decimal("7000")

COLUMN_ALIASES = {
    "номер объекта": "Номер объекта",
    "номер обьекта": "Номер объекта",
    "объект": "Номер объекта",
    "обьект": "Номер объекта",
    "код объекта": "Номер объекта",
    "код обьекта": "Номер объекта",
    "адрес": "Адрес",
    "компания": "Компания",
    "реквизиты": "Реквизиты",
    "реквезиты": "Реквизиты",
    "протокол": "Протокол",
    "ставка": "Ставка",
    "подсудность": "Подсудность",
    "суд": "Подсудность",
    "судебный участок": "Подсудность",
    "адрес суда": "Адрес суда",
}


@dataclass(frozen=True)
class CourtOrderBaseData:
    object_code: str
    object_address: str
    company: str
    requisites: str
    protocol: str
    monthly_rate: Decimal | None

    @property
    def claimant_name(self) -> str:
        match = re.search(r'ООО\s+"([^"]+)"', self.requisites)
        if match:
            return f'ООО "{match.group(1)}"'
        return self.company

    @property
    def inn(self) -> str:
        return _extract_regex(self.requisites, r"ИНН\s+(\d+)")

    @property
    def kpp(self) -> str:
        return _extract_regex(self.requisites, r"КПП\s+(\d+)")

    @property
    def ogrn(self) -> str:
        return _extract_regex(self.requisites, r"ОГРН\s+(\d+)")

    @property
    def legal_address(self) -> str:
        return _extract_between(self.requisites, "Юридический адрес", "Генеральный директор")

    @property
    def director_name(self) -> str:
        text = _extract_between(self.requisites, "Генеральный директор", "Главный бухгалтер")
        if text:
            return text
        return _extract_regex(self.requisites, r"Генеральный директор\s+(.+?)(?:\s+Расчетный счет|$)")


@dataclass(frozen=True)
class CourtJurisdiction:
    object_code: str
    court: str
    court_address: str


@dataclass(frozen=True)
class CourtOrderStaticData(CourtOrderBaseData):
    court: str
    court_address: str


@dataclass(frozen=True)
class WorkbookValidationReport:
    kind: str
    objects_count: int
    warning_counts: dict[str, int]

    @property
    def warnings_count(self) -> int:
        return sum(self.warning_counts.values())


@dataclass(frozen=True)
class GenerationIssue:
    severity: str
    account_number: str
    debtor_name: str
    object_code: str
    reason: str
    source_row: int = 0
    company: str = ""
    source_amount: Decimal | None = None
    calculated_amount: Decimal | None = None


@dataclass(frozen=True)
class CourtOrderGenerationResult:
    zip_path: str
    documents_count: int
    skipped_count: int
    warnings_count: int
    total_debt_amount: Decimal


def normalize_header(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip().lower()).replace("ё", "е")


def object_code_from_account(account_number: str) -> str:
    digits = re.sub(r"\D+", "", account_number)
    if len(digits) != 8:
        return ""
    return digits[:4]


def read_optional_columns_by_account(
    registry_path: str | Path,
    columns_by_name: dict[str, set[str]],
) -> dict[str, dict[str, Any]]:
    path = Path(registry_path)
    workbook = load_workbook(path, data_only=True)
    worksheet = workbook.active
    expected = {
        normalize_registry_header(alias): canonical
        for canonical, aliases in columns_by_name.items()
        for alias in aliases
    }

    account_col: int | None = None
    value_cols: dict[str, int] = {}
    header_row: int | None = None

    for row_idx in range(1, min(worksheet.max_row, 30) + 1):
        for col_idx in range(1, worksheet.max_column + 1):
            header = normalize_registry_header(worksheet.cell(row_idx, col_idx).value)
            if header == normalize_registry_header("Лицевой счет"):
                account_col = col_idx
            canonical = expected.get(header)
            if canonical:
                value_cols[canonical] = col_idx
        if account_col is not None and value_cols:
            header_row = row_idx
            break

    if header_row is None or account_col is None:
        return {}

    values: dict[str, dict[str, Any]] = {}
    for row_idx in range(header_row + 1, worksheet.max_row + 1):
        account = str(worksheet.cell(row_idx, account_col).value or "").strip()
        if not account:
            continue
        row_values = {
            canonical: worksheet.cell(row_idx, col_idx).value
            for canonical, col_idx in value_cols.items()
            if worksheet.cell(row_idx, col_idx).value not in (None, "")
        }
        if row_values:
            values[account] = row_values
    return values


def read_optional_money_by_account(
    registry_path: str | Path,
    column_names: set[str],
) -> dict[str, Decimal]:
    rows = read_optional_columns_by_account(
        registry_path,
        {"value": column_names},
    )
    return {
        account: parse_money(values["value"])
        for account, values in rows.items()
        if "value" in values
    }


def _extract_regex(text: str, pattern: str) -> str:
    match = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
    return " ".join(match.group(1).split()) if match else ""


def _extract_between(text: str, start: str, end: str) -> str:
    pattern = re.escape(start) + r"\s+(.+?)\s+" + re.escape(end)
    return _extract_regex(text, pattern)


def _replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return

    original = "".join(run.text for run in paragraph.runs)
    replaced = original
    for key, value in replacements.items():
        replaced = replaced.replace("{{ " + key + " }}", value)
        replaced = replaced.replace("{{" + key + "}}", value)
    replaced = re.sub(r" {2,}", " ", replaced)

    if replaced != original:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_doc(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                _replace_text_in_paragraph(paragraph, replacements)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _replace_text_in_paragraph(paragraph, replacements)


def _normalize_document_font(
    doc: Document,
    *,
    font_name: str = COURT_ORDER_FONT_NAME,
    font_size_pt: int = COURT_ORDER_FONT_SIZE_PT,
) -> None:
    roots = [doc.element]
    for section in doc.sections:
        roots.extend((section.header._element, section.footer._element))

    seen_roots: set[int] = set()
    for root in roots:
        root_id = id(root)
        if root_id in seen_roots:
            continue
        seen_roots.add(root_id)

        for run_element in root.iter(qn("w:r")):
            run = Run(run_element, None)
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            run_properties = run_element.get_or_add_rPr()
            run_fonts = run_properties.get_or_add_rFonts()
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                run_fonts.set(qn(f"w:{attribute}"), font_name)
            complex_size = run_properties.find(qn("w:szCs"))
            if complex_size is None:
                complex_size = OxmlElement("w:szCs")
                run_properties.append(complex_size)
            complex_size.set(qn("w:val"), str(font_size_pt * 2))


def _read_header_map(
    worksheet,
    required_columns: list[str],
    source_label: str,
) -> dict[str, int]:
    header_map: dict[str, int] = {}
    for col_idx in range(1, worksheet.max_column + 1):
        normalized = normalize_header(worksheet.cell(1, col_idx).value)
        if not normalized:
            continue
        canonical = COLUMN_ALIASES.get(normalized)
        if canonical:
            header_map[canonical] = col_idx

    missing = [column for column in required_columns if column not in header_map]
    if missing:
        raise ValueError(
            f"В файле «{source_label}» не найдены обязательные колонки: "
            + ", ".join(missing)
        )
    return header_map


def _load_source_worksheet(source_path: str | Path, source_label: str):
    path = Path(source_path)
    if not path.exists():
        raise FileNotFoundError(f"Файл «{source_label}» не найден: {path}")
    if not path.is_file():
        raise ValueError(f"Путь «{source_label}» должен указывать на файл .xlsx: {path}")
    if path.suffix.lower() not in {".xlsx", ".xlsm"}:
        raise ValueError(f"Файл «{source_label}» должен иметь формат .xlsx или .xlsm.")
    try:
        workbook = load_workbook(path, data_only=True)
    except Exception as exc:
        raise ValueError(f"Не удалось прочитать файл «{source_label}» как Excel.") from exc
    return workbook.active


def _parse_object_code(value: Any, row_idx: int, source_label: str) -> str:
    digits = re.sub(r"\D+", "", str(value or ""))
    if not digits:
        return ""
    if len(digits) > 4:
        raise ValueError(
            f"В файле «{source_label}» строка {row_idx}: код объекта должен состоять из 4 цифр."
        )
    return digits.zfill(4)


def load_base_data(base_data_path: str | Path) -> dict[str, CourtOrderBaseData]:
    source_label = "Основная БЗ"
    worksheet = _load_source_worksheet(base_data_path, source_label)
    header_map = _read_header_map(worksheet, REQUIRED_BASE_COLUMNS, source_label)
    by_object: dict[str, CourtOrderBaseData] = {}
    row_by_object: dict[str, int] = {}

    for row_idx in range(2, worksheet.max_row + 1):
        values = {
            column: worksheet.cell(row_idx, col_idx).value
            for column, col_idx in header_map.items()
        }
        object_code = _parse_object_code(values.get("Номер объекта"), row_idx, source_label)
        if not object_code:
            continue
        if object_code in by_object:
            raise ValueError(
                f"В файле «{source_label}» повторяется код объекта {object_code}: "
                f"строки {row_by_object[object_code]} и {row_idx}."
            )

        company = str(values.get("Компания") or "").strip()
        rate = None
        if values.get("Ставка") not in (None, ""):
            try:
                rate = parse_money(values["Ставка"])
            except ValueError as exc:
                raise ValueError(
                    f"В файле «{source_label}» строка {row_idx}: некорректная ставка."
                ) from exc

        by_object[object_code] = CourtOrderBaseData(
            object_code=object_code,
            object_address=str(values.get("Адрес") or "").strip(),
            company=company,
            requisites=str(values.get("Реквизиты") or "").strip(),
            protocol=str(values.get("Протокол") or "").strip(),
            monthly_rate=rate,
        )
        row_by_object[object_code] = row_idx

    if not by_object:
        raise ValueError("В основной БЗ нет строк с кодами объектов.")
    return by_object


def load_jurisdiction_data(
    jurisdiction_path: str | Path,
) -> dict[str, CourtJurisdiction]:
    source_label = "БЗ подсудности"
    worksheet = _load_source_worksheet(jurisdiction_path, source_label)
    header_map = _read_header_map(
        worksheet,
        REQUIRED_JURISDICTION_COLUMNS,
        source_label,
    )
    by_object: dict[str, CourtJurisdiction] = {}
    row_by_object: dict[str, int] = {}

    for row_idx in range(2, worksheet.max_row + 1):
        values = {
            column: worksheet.cell(row_idx, col_idx).value
            for column, col_idx in header_map.items()
        }
        object_code = _parse_object_code(values.get("Номер объекта"), row_idx, source_label)
        if not object_code:
            continue
        if object_code in by_object:
            raise ValueError(
                f"В файле «{source_label}» повторяется код объекта {object_code}: "
                f"строки {row_by_object[object_code]} и {row_idx}."
            )

        court, embedded_court_address = _split_court_and_address(
            values.get("Подсудность")
        )
        explicit_court_address = str(values.get("Адрес суда") or "").strip()
        by_object[object_code] = CourtJurisdiction(
            object_code=object_code,
            court=court,
            court_address=explicit_court_address or embedded_court_address,
        )
        row_by_object[object_code] = row_idx

    if not by_object:
        raise ValueError("В БЗ подсудности нет строк с кодами объектов.")
    return by_object


def _split_court_and_address(value: Any) -> tuple[str, str]:
    lines = [
        line.strip()
        for line in str(value or "").splitlines()
        if line.strip()
    ]
    if not lines:
        return "", ""
    return lines[0], "\n".join(lines[1:])


def validate_base_data(base_data_path: str | Path) -> WorkbookValidationReport:
    data = load_base_data(base_data_path)
    warning_counts = {
        "без адреса объекта": sum(not item.object_address for item in data.values()),
        "без компании": sum(not item.company for item in data.values()),
        "без реквизитов": sum(not item.requisites for item in data.values()),
        "без протокола": sum(not item.protocol for item in data.values()),
        "без ставки": sum(item.monthly_rate is None for item in data.values()),
        "не распознан ИНН": sum(not item.inn for item in data.values()),
        "не распознан КПП": sum(not item.kpp for item in data.values()),
        "не распознан ОГРН": sum(not item.ogrn for item in data.values()),
        "не распознан юридический адрес": sum(not item.legal_address for item in data.values()),
    }
    return WorkbookValidationReport(
        kind="Основная БЗ",
        objects_count=len(data),
        warning_counts={key: count for key, count in warning_counts.items() if count},
    )


def validate_jurisdiction_data(
    jurisdiction_path: str | Path,
) -> WorkbookValidationReport:
    data = load_jurisdiction_data(jurisdiction_path)
    warning_counts = {
        "без подсудности": sum(not item.court for item in data.values()),
        "без адреса суда": sum(not item.court_address for item in data.values()),
    }
    return WorkbookValidationReport(
        kind="БЗ подсудности",
        objects_count=len(data),
        warning_counts={key: count for key, count in warning_counts.items() if count},
    )


def merge_reference_data(
    base_by_object: dict[str, CourtOrderBaseData],
    jurisdiction_by_object: dict[str, CourtJurisdiction],
) -> dict[str, CourtOrderStaticData]:
    merged: dict[str, CourtOrderStaticData] = {}
    for object_code, base in base_by_object.items():
        jurisdiction = jurisdiction_by_object.get(object_code)
        merged[object_code] = CourtOrderStaticData(
            object_code=base.object_code,
            object_address=base.object_address,
            company=base.company,
            requisites=base.requisites,
            protocol=base.protocol,
            monthly_rate=base.monthly_rate,
            court=jurisdiction.court if jurisdiction else "",
            court_address=jurisdiction.court_address if jurisdiction else "",
        )
    return merged


def load_static_data(static_data_path: str | Path) -> dict[str, CourtOrderStaticData]:
    """Backward-compatible loader for an old combined workbook."""
    return merge_reference_data(
        load_base_data(static_data_path),
        load_jurisdiction_data(static_data_path),
    )


def calculate_state_duty(amount: Decimal) -> Decimal:
    amount = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    if amount <= 0:
        return Decimal("0.00")
    if amount <= Decimal("100000"):
        duty = Decimal("4000")
    elif amount <= Decimal("300000"):
        duty = Decimal("4000") + (amount - Decimal("100000")) * Decimal("0.03")
    elif amount <= Decimal("500000"):
        duty = Decimal("10000") + (amount - Decimal("300000")) * Decimal("0.025")
    elif amount <= Decimal("1000000"):
        duty = Decimal("15000") + (amount - Decimal("500000")) * Decimal("0.02")
    elif amount <= Decimal("3000000"):
        duty = Decimal("25000") + (amount - Decimal("1000000")) * Decimal("0.01")
    elif amount <= Decimal("8000000"):
        duty = Decimal("45000") + (amount - Decimal("3000000")) * Decimal("0.007")
    elif amount <= Decimal("24000000"):
        duty = Decimal("80000") + (amount - Decimal("8000000")) * Decimal("0.0035")
    elif amount <= Decimal("50000000"):
        duty = Decimal("136000") + (amount - Decimal("24000000")) * Decimal("0.003")
    elif amount <= Decimal("100000000"):
        duty = Decimal("214000") + (amount - Decimal("50000000")) * Decimal("0.002")
    else:
        duty = min(
            Decimal("314000") + (amount - Decimal("100000000")) * Decimal("0.0015"),
            Decimal("900000"),
        )
    return (duty / Decimal("2")).quantize(Decimal("1"), rounding=ROUND_FLOOR)


def calculate_court_debt(
    source_debt: Decimal,
    monthly_rate: Decimal,
) -> tuple[int, Decimal]:
    if monthly_rate <= 0 or monthly_rate != monthly_rate.to_integral_value():
        raise ValueError("Ставка должна быть положительным целым числом.")
    full_months = int((source_debt / monthly_rate).to_integral_value(rounding=ROUND_FLOOR))
    claimable_months = min(full_months, LIMITATION_PERIOD_MONTHS)
    return claimable_months, monthly_rate * claimable_months


def calculate_court_period(period_text: str, full_months: int) -> str:
    dates = [
        datetime.strptime(value, "%d.%m.%Y")
        for value in re.findall(r"\d{2}\.\d{2}\.\d{4}", period_text)
    ]
    if not dates:
        raise ValueError("Не удалось определить последний месяц периода задолженности.")
    end = max(dates).replace(day=1)
    start_month_index = end.year * 12 + end.month - 1 - (full_months - 1)
    start = datetime(start_month_index // 12, start_month_index % 12 + 1, 1)
    return f"{start:%d.%m.%Y} - {end:%d.%m.%Y}"


def court_order_exclusion_reason(full_months: int, court_debt: Decimal) -> str | None:
    if full_months == 0:
        return "Задолженность меньше одной месячной ставки."
    if court_debt < MIN_COURT_ORDER_AMOUNT:
        return "Сумма для судебного приказа менее 7 000 рублей."
    if court_debt > Decimal("500000"):
        return "Сумма превышает предел приказного производства — 500 000 рублей."
    return None


def build_replacements(
    row: ClaimRow,
    static_data: CourtOrderStaticData,
    application_date: str,
    penalty_amount: Decimal | None = None,
    debtor_data: dict[str, Any] | None = None,
) -> dict[str, str]:
    debtor_data = debtor_data or {}
    total_claim_amount = row.debt_amount
    state_duty = calculate_state_duty(total_claim_amount)
    context = {
        "court_name": static_data.court,
        "court_address": static_data.court_address,
        "claimant_name": static_data.claimant_name,
        "claimant_company": static_data.company,
        "claimant_requisites": static_data.requisites,
        "claimant_ogrn": static_data.ogrn,
        "claimant_inn": static_data.inn,
        "claimant_kpp": static_data.kpp,
        "claimant_legal_address": static_data.legal_address,
        "claimant_director_name": static_data.director_name,
        "debtor_name": row.debtor_name,
        "debtor_birth_date": str(debtor_data.get("Дата рождения") or ""),
        "debtor_birth_place": str(debtor_data.get("Место рождения") or ""),
        "debtor_registration_address": str(debtor_data.get("Адрес регистрации") or ""),
        "debtor_passport": str(debtor_data.get("Идентификатор должника") or ""),
        "account_number": row.account_number,
        "object_code": static_data.object_code,
        "object_address": static_data.object_address or row.object_address,
        "parking_place_number": row.parking_place_number,
        "debt_period": row.debt_period,
        "debt_amount": format_money(row.debt_amount),
        "debt_amount_words": money_to_words(row.debt_amount),
        "penalty_amount": "__________",
        "penalty_amount_words": "",
        "total_claim_amount": format_money(total_claim_amount),
        "total_claim_amount_words": money_to_words(total_claim_amount),
        "state_duty": format_money(state_duty),
        "state_duty_words": money_to_words(state_duty),
        "monthly_rate": format_money(static_data.monthly_rate) if static_data.monthly_rate is not None else "",
        "protocol": static_data.protocol,
        "application_date": application_date,
    }
    return {key: str(value) for key, value in context.items()}


def _project_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _base_data_errors(static_data: CourtOrderStaticData) -> list[str]:
    missing: list[str] = []
    if not static_data.object_address:
        missing.append("адрес объекта")
    if not static_data.company:
        missing.append("компания")
    if not static_data.requisites:
        missing.append("реквизиты")
    if not static_data.protocol:
        missing.append("протокол")
    if (
        static_data.monthly_rate is None
        or static_data.monthly_rate <= 0
        or static_data.monthly_rate != static_data.monthly_rate.to_integral_value()
    ):
        missing.append("ставка")
    if not static_data.inn:
        missing.append("ИНН")
    if not static_data.kpp:
        missing.append("КПП")
    if not static_data.ogrn:
        missing.append("ОГРН")
    if not static_data.legal_address:
        missing.append("юридический адрес")
    return missing


def _debtor_data_warning(debtor_data: dict[str, Any]) -> str | None:
    required = [
        "Дата рождения",
        "Место рождения",
        "Адрес регистрации",
        "Идентификатор должника",
    ]
    missing = [field for field in required if not debtor_data.get(field)]
    if not missing:
        return None
    return "Не заполнены данные должника: " + ", ".join(missing)


def _write_issues_workbook(
    issues: list[GenerationIssue],
    output_path: Path,
    *,
    documents_count: int,
    total_amount: Decimal,
) -> Path:
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Итоги"
    summary.append(["Показатель", "Значение"])
    summary.append(["Создано заявлений", documents_count])
    summary.append(["Пропущено строк", sum(issue.severity != "warning" for issue in issues)])
    summary.append(["Ошибок", sum(issue.severity == "error" for issue in issues)])
    summary.append(["Предупреждений", sum(issue.severity == "warning" for issue in issues)])
    summary.append(["Не сформировано", sum(issue.severity == "not_generated" for issue in issues)])
    summary.append(["Добросовестных плательщиков", sum(issue.severity == "good_payer" for issue in issues)])
    summary.append(["Дублей", sum(issue.severity == "duplicate" for issue in issues)])
    summary.append(["Сумма сформированных заявлений", total_amount])
    for cell in summary[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
    summary.column_dimensions["A"].width = 40
    summary.column_dimensions["B"].width = 22
    summary["B9"].number_format = "#,##0.00"
    headers = [
        "Строка реестра",
        "Лицевой счет",
        "ФИО",
        "Код объекта",
        "Компания",
        "Исходная задолженность",
        "Рассчитанная задолженность",
        "Причина",
    ]

    for severity, title in (
        ("error", "Ошибки"),
        ("warning", "Предупреждения"),
        ("not_generated", "Не сформированы"),
        ("good_payer", "Добросовестные плательщики"),
        ("duplicate", "Дубли"),
    ):
        rows = [issue for issue in issues if issue.severity == severity]
        if not rows:
            continue
        worksheet = workbook.create_sheet(title)
        worksheet.append(headers)
        for issue in rows:
            worksheet.append(
                [
                    issue.source_row or "",
                    issue.account_number,
                    issue.debtor_name,
                    issue.object_code,
                    issue.company,
                    issue.source_amount,
                    issue.calculated_amount,
                    issue.reason,
                ]
            )
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = f"A1:H{worksheet.max_row}"
        for cell in worksheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E78")
        worksheet.column_dimensions["A"].width = 16
        worksheet.column_dimensions["B"].width = 18
        worksheet.column_dimensions["C"].width = 32
        worksheet.column_dimensions["D"].width = 16
        worksheet.column_dimensions["E"].width = 32
        worksheet.column_dimensions["F"].width = 22
        worksheet.column_dimensions["G"].width = 24
        worksheet.column_dimensions["H"].width = 70
        for column in ("F", "G"):
            for cell in worksheet[column][1:]:
                cell.number_format = "#,##0.00"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)
    return output_path


def generate_court_orders_zip_result(
    registry_path: str,
    template_path: str,
    static_data_path: str | None = None,
    output_zip_path: str | None = None,
    application_date: str | None = None,
    *,
    base_data_path: str | None = None,
    jurisdiction_path: str | None = None,
) -> CourtOrderGenerationResult:
    base_data_path = base_data_path or static_data_path
    jurisdiction_path = jurisdiction_path or static_data_path
    if not base_data_path or not jurisdiction_path:
        raise ValueError("Укажите пути к основной БЗ и БЗ подсудности.")
    if not output_zip_path:
        raise ValueError("Не указан путь к итоговому ZIP-архиву.")

    template = Path(template_path)
    output_zip = Path(output_zip_path)
    if not template.exists():
        raise FileNotFoundError(f"Word-шаблон судебного заявления не найден: {template}")

    registry_issues: list[ClaimIssue] = []
    rows = read_registry(registry_path, issues=registry_issues)
    if not rows and not registry_issues:
        raise ValueError("В реестре нет строк с положительной суммой долга для генерации заявлений.")

    base_by_object = load_base_data(base_data_path)
    jurisdiction_by_object = load_jurisdiction_data(jurisdiction_path)
    static_by_object = merge_reference_data(base_by_object, jurisdiction_by_object)
    debtor_data_by_account = read_optional_columns_by_account(
        registry_path,
        {
            "Дата рождения": {"Дата рождения", "Дата рождения должника"},
            "Место рождения": {"Место рождения", "Место рождения должника"},
            "Адрес регистрации": {
                "Адрес регистрации",
                "Адрес регистрации должника",
            },
            "Идентификатор должника": {
                "Идентификатор должника",
                "Паспорт",
                "Паспортные данные",
                "СНИЛС",
                "ИНН должника",
            },
        },
    )

    temp_root = _project_root() / "storage" / "temp"
    temp_root.mkdir(parents=True, exist_ok=True)
    temp_dir = temp_root / f"court_orders_{uuid.uuid4().hex}"
    temp_dir.mkdir(parents=True, exist_ok=False)

    generated_files: list[Path] = []
    issues: list[GenerationIssue] = [
        GenerationIssue(
            severity=issue.category,
            account_number=issue.account_number,
            debtor_name=issue.debtor_name,
            object_code=object_code_from_account(issue.account_number),
            reason=issue.reason,
            source_row=issue.source_row,
            company=issue.company,
            source_amount=issue.debt_amount,
        )
        for issue in registry_issues
    ]
    used_names: set[str] = set()
    skipped_count = len(registry_issues)
    total_debt_amount = Decimal("0.00")
    effective_date = application_date or date.today().strftime("%d.%m.%Y")

    try:
        for index, row in enumerate(rows, start=1):
            object_code = object_code_from_account(row.account_number)
            if not object_code:
                issues.append(
                    GenerationIssue(
                        severity="error",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code="",
                        reason="Не удалось определить код объекта из лицевого счета.",
                        source_row=row.source_row,
                        company=row.company,
                        source_amount=row.debt_amount,
                    )
                )
                skipped_count += 1
                continue

            static_data = static_by_object.get(object_code)
            if static_data is None:
                issues.append(
                    GenerationIssue(
                        severity="error",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason="Код объекта отсутствует в основной БЗ.",
                        source_row=row.source_row,
                        company=row.company,
                        source_amount=row.debt_amount,
                    )
                )
                skipped_count += 1
                continue
            missing_base_fields = _base_data_errors(static_data)
            if missing_base_fields:
                issues.append(
                    GenerationIssue(
                        severity="error",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason="В основной БЗ не заполнены: " + ", ".join(missing_base_fields),
                        source_row=row.source_row,
                        company=static_data.company or row.company,
                        source_amount=row.debt_amount,
                    )
                )
                skipped_count += 1
                continue

            full_months, court_debt = calculate_court_debt(
                row.debt_amount,
                static_data.monthly_rate,
            )
            limitation_cap = static_data.monthly_rate * LIMITATION_PERIOD_MONTHS
            if row.debt_amount > limitation_cap:
                issues.append(
                    GenerationIssue(
                        severity="warning",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason=(
                            "Сумма задолженности ограничена сроком исковой "
                            f"давности {LIMITATION_PERIOD_MONTHS} месяцев."
                        ),
                        source_row=row.source_row,
                        company=static_data.company,
                        source_amount=row.debt_amount,
                        calculated_amount=court_debt,
                    )
                )
            exclusion_reason = court_order_exclusion_reason(full_months, court_debt)
            if exclusion_reason:
                issues.append(
                    GenerationIssue(
                        severity="not_generated",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason=exclusion_reason,
                        source_row=row.source_row,
                        company=static_data.company,
                        source_amount=row.debt_amount,
                        calculated_amount=court_debt,
                    )
                )
                skipped_count += 1
                continue
            try:
                court_period = calculate_court_period(row.debt_period, full_months)
            except ValueError as exc:
                issues.append(
                    GenerationIssue(
                        severity="error",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason=str(exc),
                        source_row=row.source_row,
                        company=static_data.company,
                        source_amount=row.debt_amount,
                        calculated_amount=court_debt,
                    )
                )
                skipped_count += 1
                continue

            court_row = replace(
                row,
                debt_amount=court_debt,
                debt_period=court_period,
            )
            doc = Document(str(template))
            debtor_data = debtor_data_by_account.get(row.account_number, {})
            warning = _debtor_data_warning(debtor_data)
            if warning:
                issues.append(
                    GenerationIssue(
                        severity="warning",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason=warning,
                        source_row=row.source_row,
                        company=static_data.company,
                        source_amount=row.debt_amount,
                        calculated_amount=court_debt,
                    )
                )
            if not static_data.court:
                issues.append(
                    GenerationIssue(
                        severity="warning",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason="В БЗ подсудности не заполнен судебный участок.",
                        source_row=row.source_row,
                        company=static_data.company,
                        source_amount=row.debt_amount,
                        calculated_amount=court_debt,
                    )
                )
            if not static_data.court_address:
                issues.append(
                    GenerationIssue(
                        severity="warning",
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        object_code=object_code,
                        reason="В БЗ подсудности не заполнен адрес суда.",
                        source_row=row.source_row,
                        company=static_data.company,
                        source_amount=row.debt_amount,
                        calculated_amount=court_debt,
                    )
                )
            _replace_in_doc(
                doc,
                build_replacements(
                    court_row,
                    static_data,
                    effective_date,
                    None,
                    debtor_data,
                ),
            )
            _normalize_document_font(doc)

            base_name = safe_filename(f"{index:03d}_{row.account_number}_{row.debtor_name}_судебный_приказ")
            file_name = base_name + ".docx"
            counter = 2
            while file_name.lower() in used_names:
                file_name = f"{base_name}_{counter}.docx"
                counter += 1
            used_names.add(file_name.lower())

            company_dir = temp_dir / safe_filename(static_data.company)
            company_dir.mkdir(parents=True, exist_ok=True)
            file_path = company_dir / file_name
            doc.save(file_path)
            generated_files.append(file_path)
            total_debt_amount += court_debt

        if not generated_files and not issues:
            raise ValueError(
                "Не создано ни одного заявления и не сформирован отчет об ошибках."
            )

        if issues:
            generated_files.append(
                _write_issues_workbook(
                    issues,
                    temp_dir / "errors.xlsx",
                    documents_count=sum(
                        path.suffix.lower() == ".docx"
                        for path in generated_files
                    ),
                    total_amount=total_debt_amount,
                )
            )
        create_zip(generated_files, output_zip, base_dir=temp_dir)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

    return CourtOrderGenerationResult(
        zip_path=str(output_zip),
        documents_count=sum(path.suffix.lower() == ".docx" for path in generated_files),
        skipped_count=skipped_count,
        warnings_count=sum(issue.severity == "warning" for issue in issues),
        total_debt_amount=total_debt_amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP),
    )


def generate_court_orders_zip(
    registry_path: str,
    template_path: str,
    static_data_path: str | None = None,
    output_zip_path: str | None = None,
    application_date: str | None = None,
    *,
    base_data_path: str | None = None,
    jurisdiction_path: str | None = None,
) -> str:
    return generate_court_orders_zip_result(
        registry_path=registry_path,
        template_path=template_path,
        static_data_path=static_data_path,
        output_zip_path=output_zip_path,
        application_date=application_date,
        base_data_path=base_data_path,
        jurisdiction_path=jurisdiction_path,
    ).zip_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Генерация Word-заявлений о вынесении судебного приказа из Excel-реестра"
    )
    parser.add_argument("registry", help="Готовый Excel-реестр")
    parser.add_argument("--out", default="storage/output/court_orders.zip", help="Путь к итоговому ZIP-архиву")
    parser.add_argument(
        "--template",
        default="app/modules/court_orders/court_order_template.docx",
        help="Word-шаблон заявления",
    )
    parser.add_argument(
        "--base-data",
        default="storage/court_orders/base_data.xlsx",
        help="Основная БЗ объектов, компаний и реквизитов",
    )
    parser.add_argument(
        "--jurisdiction",
        default="storage/court_orders/jurisdiction.xlsx",
        help="БЗ подсудности",
    )
    parser.add_argument("--application-date", default=None, help="Дата заявления, по умолчанию сегодня")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    try:
        result = generate_court_orders_zip_result(
            registry_path=args.registry,
            template_path=args.template,
            base_data_path=args.base_data,
            jurisdiction_path=args.jurisdiction,
            output_zip_path=args.out,
            application_date=args.application_date,
        )
    except (FileNotFoundError, ValueError) as exc:
        print(f"Ошибка: {exc}", file=sys.stderr)
        raise SystemExit(1) from exc

    print(f"Создано заявлений: {result.documents_count}")
    if result.skipped_count:
        print(f"Пропущено строк: {result.skipped_count}")
    if result.warnings_count:
        print(f"Предупреждений: {result.warnings_count}")
    print(f"Итоговая сумма долга: {format_money(result.total_debt_amount)}")
    print(f"Архив: {result.zip_path}")


if __name__ == "__main__":
    main()
