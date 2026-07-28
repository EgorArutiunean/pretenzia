from __future__ import annotations

import argparse
import re
import shutil
import sys
import uuid
from dataclasses import dataclass, field
from datetime import date, timedelta
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill

from app.shared.file_utils import safe_filename
from app.shared.money_to_words import money_to_words
from app.shared.object_data import ObjectData, load_object_data
from app.shared.zip_utils import create_zip


REQUIRED_COLUMNS = [
    "Лицевой счет",
    "ФИО",
    "Адрес",
    "Период задолженности",
    "Сумма долга",
]

DEFAULT_CONTEXT = {
    "company_name": "ООО «ПАРКИНГ ЛАЙН»",
    "company_ogrn": "1207700184977",
    "company_inn": "7727444957",
    "company_kpp": "772701001",
    "company_post_address": "117042, город Москва, Чечёрский проезд, д. 120, помещ. 1/1",
    "company_email": "_______",
    "director_position": "Генеральный директор",
    "director_name": "Голубев И. А.",
}


@dataclass(frozen=True)
class ClaimRow:
    account_number: str
    debtor_name: str
    address: str
    debt_period: str
    debt_amount: Decimal
    company: str = ""
    company_inn: str = ""
    company_code: str = ""
    source_row: int = field(default=0, compare=False)

    @property
    def parking_place_number(self) -> str:
        digits = re.sub(r"\D+", "", self.account_number)
        return digits[-4:] if len(digits) >= 4 else digits

    @property
    def object_address(self) -> str:
        return re.sub(
            r",?\s*машиноместо\s*№\s*\d+\s*$",
            "",
            self.address,
            flags=re.IGNORECASE,
        ).strip()


@dataclass(frozen=True)
class ClaimsGenerationResult:
    zip_path: str
    documents_count: int
    total_amount: Decimal
    skipped_count: int = 0
    errors_count: int = 0


@dataclass(frozen=True)
class ClaimIssue:
    category: str
    source_row: int
    account_number: str
    debtor_name: str
    company: str
    debt_amount: Decimal | None
    reason: str


def parse_money(value: Any) -> Decimal:
    if value is None:
        return Decimal("0.00")
    if isinstance(value, Decimal):
        return value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    if isinstance(value, (int, float)):
        return Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

    text = str(value).strip().replace("\u00a0", " ")
    text = re.sub(r"[^0-9,\.\- ]+", "", text).replace(" ", "")
    if "," in text and "." in text:
        if text.rfind(",") > text.rfind("."):
            text = text.replace(".", "").replace(",", ".")
        else:
            text = text.replace(",", "")
    else:
        text = text.replace(",", ".")

    try:
        return Decimal(text).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except (InvalidOperation, ValueError) as exc:
        raise ValueError(f"Не удалось распознать сумму долга: {value!r}") from exc


def format_money(value: Decimal) -> str:
    value = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    integer, frac = f"{value:.2f}".split(".")
    return f"{int(integer):,}".replace(",", " ") + f",{frac}"


def normalize_header(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def read_registry(
    registry_path: str | Path,
    sheet_name: str | None = None,
    *,
    issues: list[ClaimIssue] | None = None,
) -> list[ClaimRow]:
    path = Path(registry_path)
    if not path.exists():
        raise FileNotFoundError(f"Excel-реестр не найден: {path}")

    wb = load_workbook(path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active

    header_row: int | None = None
    columns: dict[str, int] = {}
    required_by_norm = {normalize_header(col): col for col in REQUIRED_COLUMNS}

    for row_idx in range(1, min(ws.max_row, 30) + 1):
        found: dict[str, int] = {}
        for col_idx in range(1, ws.max_column + 1):
            norm = normalize_header(ws.cell(row_idx, col_idx).value)
            if norm in required_by_norm:
                found[required_by_norm[norm]] = col_idx
        if len(found) == len(REQUIRED_COLUMNS):
            header_row = row_idx
            columns = found
            break

    if header_row is None:
        visible_headers = [
            str(ws.cell(1, col_idx).value)
            for col_idx in range(1, ws.max_column + 1)
            if ws.cell(1, col_idx).value not in (None, "")
        ]
        raise ValueError(
            "В Excel-реестре не найдены обязательные колонки: "
            + ", ".join(REQUIRED_COLUMNS)
            + ". Найденные заголовки первой строки: "
            + (", ".join(visible_headers) if visible_headers else "нет")
        )

    optional_columns = {
        normalize_header("Компания"): "Компания",
        normalize_header("ИНН компании"): "ИНН компании",
        normalize_header("Код компании"): "Код компании",
    }
    for col_idx in range(1, ws.max_column + 1):
        canonical = optional_columns.get(normalize_header(ws.cell(header_row, col_idx).value))
        if canonical:
            columns[canonical] = col_idx

    rows: list[ClaimRow] = []
    for row_idx in range(header_row + 1, ws.max_row + 1):
        raw_values = {
            column: ws.cell(row_idx, col_idx).value
            for column, col_idx in columns.items()
        }
        if all(value in (None, "") for value in raw_values.values()):
            continue

        try:
            amount = parse_money(raw_values["Сумма долга"])
        except ValueError as exc:
            if issues is None:
                raise
            issues.append(
                ClaimIssue(
                    category="error",
                    source_row=row_idx,
                    account_number=str(raw_values["Лицевой счет"] or "").strip(),
                    debtor_name=str(raw_values["ФИО"] or "").strip(),
                    company=str(raw_values.get("Компания") or "").strip(),
                    debt_amount=None,
                    reason=str(exc),
                )
            )
            continue
        if amount <= 0:
            if issues is not None:
                issues.append(
                    ClaimIssue(
                        category="good_payer",
                        source_row=row_idx,
                        account_number=str(raw_values["Лицевой счет"] or "").strip(),
                        debtor_name=str(raw_values["ФИО"] or "").strip(),
                        company=str(raw_values.get("Компания") or "").strip(),
                        debt_amount=amount,
                        reason="Задолженность меньше или равна нулю.",
                    )
                )
            continue

        account_number = str(raw_values["Лицевой счет"] or "").strip()
        debtor_name = str(raw_values["ФИО"] or "").strip()
        address = str(raw_values["Адрес"] or "").strip()
        debt_period = str(raw_values["Период задолженности"] or "").strip()
        if not re.fullmatch(r"\d{8}", account_number):
            reason = f"Строка {row_idx}: лицевой счёт должен состоять ровно из 8 цифр."
            if issues is None:
                raise ValueError(reason)
            issues.append(
                ClaimIssue(
                    category="error",
                    source_row=row_idx,
                    account_number=account_number,
                    debtor_name=debtor_name,
                    company=str(raw_values.get("Компания") or "").strip(),
                    debt_amount=amount,
                    reason=reason,
                )
            )
            continue
        missing = [
            label
            for label, value in (
                ("ФИО", debtor_name),
                ("Адрес", address),
                ("Период задолженности", debt_period),
            )
            if not value
        ]
        if missing:
            reason = (
                f"Строка {row_idx}: не заполнены обязательные поля: "
                + ", ".join(missing)
            )
            if issues is None:
                raise ValueError(reason)
            issues.append(
                ClaimIssue(
                    category="error",
                    source_row=row_idx,
                    account_number=account_number,
                    debtor_name=debtor_name,
                    company=str(raw_values.get("Компания") or "").strip(),
                    debt_amount=amount,
                    reason=reason,
                )
            )
            continue

        candidate = ClaimRow(
            account_number=account_number,
            debtor_name=debtor_name,
            address=address,
            debt_period=re.sub(r"\s+-\s+", " - ", debt_period),
            debt_amount=amount,
            company=str(raw_values.get("Компания") or "").strip(),
            company_inn=str(raw_values.get("ИНН компании") or "").strip(),
            company_code=str(raw_values.get("Код компании") or "").strip(),
            source_row=row_idx,
        )
        if candidate in rows:
            if issues is not None:
                issues.append(
                    ClaimIssue(
                        category="duplicate",
                        source_row=row_idx,
                        account_number=account_number,
                        debtor_name=debtor_name,
                        company=candidate.company,
                        debt_amount=amount,
                        reason="Полностью одинаковая строка уже была обработана.",
                    )
                )
            continue
        rows.append(candidate)

    return rows


def _replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return

    original = "".join(run.text for run in paragraph.runs)
    replaced = original
    for key, value in replacements.items():
        replaced = replaced.replace("{{ " + key + " }}", value)
        replaced = replaced.replace("{{" + key + "}}", value)
    replaced = re.sub(r" {2,}", " ", replaced)

    if replaced != original:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_doc(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                _replace_text_in_paragraph(paragraph, replacements)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _replace_text_in_paragraph(paragraph, replacements)


def _remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _remove_deprecated_claim_elements(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if "Приложение:" in paragraph.text:
            _remove_paragraph(paragraph)

    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in footer.paragraphs:
                paragraph.text = ""
            for table in list(footer.tables):
                table._element.getparent().remove(table._element)


def build_replacements(
    row: ClaimRow,
    claim_date: str,
    payment_deadline: str,
    object_data: ObjectData | None = None,
) -> dict[str, str]:
    context = dict(DEFAULT_CONTEXT)
    if object_data is not None:
        context.update(
            {
                "company_name": object_data.company,
                "company_ogrn": object_data.ogrn,
                "company_inn": object_data.inn,
                "company_kpp": object_data.kpp,
                "company_post_address": object_data.legal_address,
                "company_email": "",
                "director_name": object_data.director_name,
            }
        )
    elif row.company:
        context.update(
            {
                "company_name": row.company,
                "company_inn": row.company_inn,
            }
        )
    context.update(
        {
            "account_number": row.account_number,
            "debtor_name": row.debtor_name,
            "address": row.address,
            "debt_period": row.debt_period,
            "debt_amount": format_money(row.debt_amount),
            "debt_amount_words": money_to_words(row.debt_amount),
            "parking_place_number": row.parking_place_number,
            "object_address": row.object_address,
            "claim_date": claim_date,
            "payment_deadline": payment_deadline,
        }
    )
    return {key: str(value) for key, value in context.items()}


def _project_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _write_claim_issues(issues: list[ClaimIssue], output_path: Path) -> Path:
    workbook = Workbook()
    workbook.remove(workbook.active)
    headers = [
        "Строка реестра",
        "Лицевой счет",
        "ФИО",
        "Компания",
        "Сумма долга",
        "Причина",
    ]
    sheet_names = {
        "error": "Ошибки",
        "good_payer": "Добросовестные плательщики",
        "duplicate": "Дубли",
    }
    for category, sheet_name in sheet_names.items():
        category_issues = [issue for issue in issues if issue.category == category]
        if not category_issues:
            continue
        worksheet = workbook.create_sheet(sheet_name)
        worksheet.append(headers)
        for issue in category_issues:
            worksheet.append(
                [
                    issue.source_row,
                    issue.account_number,
                    issue.debtor_name,
                    issue.company,
                    issue.debt_amount,
                    issue.reason,
                ]
            )
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = f"A1:F{worksheet.max_row}"
        for cell in worksheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E78")
        for col_idx, width in enumerate((16, 18, 32, 32, 18, 70), start=1):
            worksheet.column_dimensions[chr(64 + col_idx)].width = width
        for cell in worksheet["E"][1:]:
            cell.number_format = "#,##0.00"
    workbook.save(output_path)
    return output_path


def generate_claims_zip_result(
    registry_path: str,
    template_path: str,
    output_zip_path: str,
    claim_date: str,
    payment_deadline: str,
    *,
    base_data_path: str | None = None,
) -> ClaimsGenerationResult:
    registry = Path(registry_path)
    template = Path(template_path)
    output_zip = Path(output_zip_path)

    if not template.exists():
        raise FileNotFoundError(f"Word-шаблон претензии не найден: {template}")

    issues: list[ClaimIssue] = []
    rows = read_registry(registry, issues=issues)
    if not rows and not issues:
        raise ValueError("В реестре нет строк с положительной суммой долга для генерации претензий.")
    object_data_by_code = load_object_data(base_data_path) if base_data_path else {}

    temp_root = _project_root() / "storage" / "temp"
    temp_root.mkdir(parents=True, exist_ok=True)
    temp_dir = temp_root / f"claims_{uuid.uuid4().hex}"
    temp_dir.mkdir(parents=True, exist_ok=False)

    try:
        generated_files: list[Path] = []
        used_names: set[str] = set()

        for index, row in enumerate(rows, start=1):
            object_code = row.account_number[:4]
            object_data = object_data_by_code.get(object_code)
            if object_data_by_code and object_data is None:
                issues.append(
                    ClaimIssue(
                        category="error",
                        source_row=row.source_row,
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        company=row.company,
                        debt_amount=row.debt_amount,
                        reason=f"Код объекта {object_code} отсутствует в основной БЗ.",
                    )
                )
                continue
            if object_data is not None and row.company:
                if normalize_header(object_data.company) != normalize_header(row.company):
                    issues.append(
                        ClaimIssue(
                            category="error",
                            source_row=row.source_row,
                            account_number=row.account_number,
                            debtor_name=row.debtor_name,
                            company=row.company,
                            debt_amount=row.debt_amount,
                            reason="Компания в реестре не совпадает с основной БЗ.",
                        )
                    )
                    continue

            try:
                doc = Document(str(template))
                _replace_in_doc(
                    doc,
                    build_replacements(row, claim_date, payment_deadline, object_data),
                )
                _remove_deprecated_claim_elements(doc)
            except Exception:
                issues.append(
                    ClaimIssue(
                        category="error",
                        source_row=row.source_row,
                        account_number=row.account_number,
                        debtor_name=row.debtor_name,
                        company=row.company,
                        debt_amount=row.debt_amount,
                        reason="Не удалось сформировать Word-документ.",
                    )
                )
                continue

            base_name = safe_filename(f"{index:03d}_{row.account_number}_{row.debtor_name}")
            file_name = base_name + ".docx"
            company_name = (
                object_data.company
                if object_data is not None
                else row.company or DEFAULT_CONTEXT["company_name"]
            )
            company_dir = temp_dir / safe_filename(company_name)
            company_dir.mkdir(parents=True, exist_ok=True)
            archive_name = f"{company_dir.name}/{file_name}".lower()
            counter = 2
            while archive_name in used_names:
                file_name = f"{base_name}_{counter}.docx"
                archive_name = f"{company_dir.name}/{file_name}".lower()
                counter += 1
            used_names.add(archive_name)

            file_path = company_dir / file_name
            doc.save(file_path)
            generated_files.append(file_path)

        if issues:
            generated_files.append(
                _write_claim_issues(issues, temp_dir / "errors.xlsx")
            )
        create_zip(generated_files, output_zip, base_dir=temp_dir)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

    return ClaimsGenerationResult(
        zip_path=str(output_zip),
        documents_count=sum(path.suffix.lower() == ".docx" for path in generated_files),
        total_amount=sum(
            (
                row.debt_amount
                for row in rows
                if not any(
                    issue.category == "error"
                    and issue.source_row == row.source_row
                    for issue in issues
                )
            ),
            Decimal("0.00"),
        ).quantize(Decimal("0.01")),
        skipped_count=len(issues),
        errors_count=sum(issue.category == "error" for issue in issues),
    )


def generate_claims_zip(
    registry_path: str,
    template_path: str,
    output_zip_path: str,
    claim_date: str,
    payment_deadline: str,
    *,
    base_data_path: str | None = None,
) -> str:
    return generate_claims_zip_result(
        registry_path=registry_path,
        template_path=template_path,
        output_zip_path=output_zip_path,
        claim_date=claim_date,
        payment_deadline=payment_deadline,
        base_data_path=base_data_path,
    ).zip_path


def parse_args() -> argparse.Namespace:
    today = date.today()
    default_deadline = today + timedelta(days=30)

    parser = argparse.ArgumentParser(description="Генерация Word-претензий из Excel-реестра и упаковка в ZIP")
    parser.add_argument("registry", help="Готовый Excel-реестр")
    parser.add_argument("--out", default="storage/output/claims.zip", help="Путь к итоговому ZIP-архиву")
    parser.add_argument("--template", default="app/modules/claims/claim_template.docx", help="Word-шаблон претензии")
    parser.add_argument("--claim-date", default=today.strftime("%d.%m.%Y"), help="Дата претензии")
    parser.add_argument("--payment-deadline", default=default_deadline.strftime("%d.%m.%Y"), help="Срок оплаты")
    parser.add_argument("--base-data", default=None, help="Основная БЗ объектов и компаний")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    try:
        result = generate_claims_zip_result(
            registry_path=args.registry,
            template_path=args.template,
            output_zip_path=args.out,
            claim_date=args.claim_date,
            payment_deadline=args.payment_deadline,
            base_data_path=args.base_data,
        )
    except (FileNotFoundError, ValueError) as exc:
        print(f"Ошибка: {exc}", file=sys.stderr)
        raise SystemExit(1) from exc

    print(f"Создано претензий: {result.documents_count}")
    print(f"Итоговая сумма долга: {format_money(result.total_amount)}")
    print(f"Архив: {result.zip_path}")


if __name__ == "__main__":
    main()
