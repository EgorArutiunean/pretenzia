from __future__ import annotations

import tempfile
import unittest
from decimal import Decimal
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from openpyxl import Workbook

from app.modules.claims.generate_claims_from_registry import (
    generate_claims_zip_result,
    read_registry,
)


def _build_registry(path: Path, rows: list[list[object]]) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Реестр"
    worksheet.append(["Лицевой счет", "ФИО", "Адрес", "Период задолженности", "Сумма долга"])
    for row in rows:
        worksheet.append(row)
    workbook.save(path)


def _build_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("Должник: {{ debtor_name }}")
    document.add_paragraph("Адрес: {{ address }}")
    document.add_paragraph("Сумма: {{ debt_amount }} ({{ debt_amount_words }})")
    document.add_paragraph("Объект: {{ object_address }}")
    document.save(path)


class ClaimsGenerationTests(unittest.TestCase):
    def test_read_registry_parses_positive_rows_only(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            registry_path = Path(temp_dir) / "registry.xlsx"
            _build_registry(
                registry_path,
                [
                    ["12970001", "Иванов Иван", "Москва, машиноместо № 0001", "01.01.2026-01.02.2026", 1000],
                    ["12970002", "Петров Петр", "Москва, машиноместо № 0002", "01.01.2026-01.02.2026", 0],
                ],
            )

            rows = read_registry(registry_path)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].account_number, "12970001")
        self.assertEqual(rows[0].debt_amount, Decimal("1000.00"))

    def test_generate_claims_zip_result_creates_one_docx_per_registry_row(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            registry_path = temp / "registry.xlsx"
            template_path = temp / "template.docx"
            output_zip_path = temp / "claims.zip"
            _build_registry(
                registry_path,
                [
                    ["12970001", "Иванов Иван", "Москва, машиноместо № 0001", "01.01.2026-01.02.2026", 1000],
                    ["12970002", "Петров Петр", "Москва, машиноместо № 0002", "01.01.2026-01.02.2026", 2500.50],
                ],
            )
            _build_template(template_path)

            result = generate_claims_zip_result(
                registry_path=str(registry_path),
                template_path=str(template_path),
                output_zip_path=str(output_zip_path),
                claim_date="01.03.2026",
                payment_deadline="31.03.2026",
            )

            with ZipFile(output_zip_path) as archive:
                names = archive.namelist()

        self.assertEqual(result.documents_count, 2)
        self.assertEqual(result.total_amount, Decimal("3500.50"))
        self.assertEqual(len(names), 2)
        self.assertTrue(all(name.endswith(".docx") for name in names))

    def test_generate_claims_zip_result_removes_deprecated_template_elements(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            registry_path = temp / "registry.xlsx"
            template_path = temp / "template.docx"
            output_zip_path = temp / "claims.zip"
            _build_registry(
                registry_path,
                [["12970001", "Тестовый Должник", "Москва, машиноместо № 0001", "01.01.2026", 1000]],
            )

            document = Document()
            document.add_paragraph("Должник: {{ debtor_name }}")
            document.add_paragraph("Приложение: расчет задолженности.")
            document.sections[0].footer.paragraphs[0].text = 'ООО "УК «Жилищник»" · ОГРН 1077746269810'
            document.save(template_path)

            generate_claims_zip_result(
                registry_path=str(registry_path),
                template_path=str(template_path),
                output_zip_path=str(output_zip_path),
                claim_date="01.03.2026",
                payment_deadline="31.03.2026",
            )

            with ZipFile(output_zip_path) as archive:
                names = archive.namelist()
                archive.extract(names[0], temp)

            generated_doc = Document(str(temp / names[0]))
            generated_text = "\n".join(paragraph.text for paragraph in generated_doc.paragraphs)
            generated_footer_text = "\n".join(
                paragraph.text
                for section in generated_doc.sections
                for footer in (section.footer, section.first_page_footer, section.even_page_footer)
                for paragraph in footer.paragraphs
            )

        self.assertNotIn("Приложение", generated_text)
        self.assertNotIn("Жилищник", generated_footer_text)
        self.assertNotIn("1077746269810", generated_footer_text)

    def test_generate_claims_uses_company_data_and_company_folders(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            registry_path = temp / "registry.xlsx"
            base_data_path = temp / "base.xlsx"
            template_path = temp / "template.docx"
            output_zip_path = temp / "claims.zip"

            registry = Workbook()
            registry_ws = registry.active
            registry_ws.append(
                [
                    "Лицевой счет",
                    "ФИО",
                    "Адрес",
                    "Период задолженности",
                    "Сумма долга",
                    "Компания",
                    "ИНН компании",
                    "Код компании",
                ]
            )
            registry_ws.append(
                ["12970001", "Иванов", "Адрес 1", "01.01.2026", 1000, "ООО Альфа", "7700000001", "A"]
            )
            registry_ws.append(
                ["56780001", "Петров", "Адрес 2", "01.01.2026", 2000, "ООО Бета", "7700000002", "B"]
            )
            registry.save(registry_path)

            base = Workbook()
            base_ws = base.active
            base_ws.append(
                [
                    "Номер объекта",
                    "Адрес",
                    "Компания",
                    "ИНН компании",
                    "Код компании",
                    "Реквизиты",
                    "Генеральный директор",
                ]
            )
            base_ws.append(["1297", "Адрес 1", "ООО Альфа", "7700000001", "A", "", "А. А. Альфов"])
            base_ws.append(["5678", "Адрес 2", "ООО Бета", "7700000002", "B", "", "Б. Б. Бетов"])
            base.save(base_data_path)

            document = Document()
            document.add_paragraph("{{ company_name }} / {{ company_inn }} / {{ director_name }}")
            document.add_paragraph("{{ debtor_name }}")
            document.save(template_path)

            result = generate_claims_zip_result(
                registry_path=str(registry_path),
                template_path=str(template_path),
                output_zip_path=str(output_zip_path),
                claim_date="01.03.2026",
                payment_deadline="31.03.2026",
                base_data_path=str(base_data_path),
            )

            with ZipFile(output_zip_path) as archive:
                names = archive.namelist()
                for name in names:
                    archive.extract(name, temp / "extracted")

            generated_texts = [
                "\n".join(p.text for p in Document(temp / "extracted" / name).paragraphs)
                for name in names
            ]

        self.assertEqual(result.documents_count, 2)
        self.assertTrue(any(name.startswith("ООО Альфа/") for name in names))
        self.assertTrue(any(name.startswith("ООО Бета/") for name in names))
        self.assertTrue(any("А. А. Альфов" in text for text in generated_texts))
        self.assertTrue(any("Б. Б. Бетов" in text for text in generated_texts))

    def test_generate_claims_zip_result_rejects_empty_registry(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            registry_path = temp / "registry.xlsx"
            template_path = temp / "template.docx"
            _build_registry(registry_path, [])
            _build_template(template_path)

            with self.assertRaisesRegex(ValueError, "нет строк"):
                generate_claims_zip_result(
                    registry_path=str(registry_path),
                    template_path=str(template_path),
                    output_zip_path=str(temp / "claims.zip"),
                    claim_date="01.03.2026",
                    payment_deadline="31.03.2026",
                )


if __name__ == "__main__":
    unittest.main()
