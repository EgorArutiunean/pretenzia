from __future__ import annotations

import tempfile
import unittest
from decimal import Decimal
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from openpyxl import Workbook, load_workbook

from app.modules.court_orders.generate_court_orders_from_registry import (
    calculate_court_debt,
    calculate_court_period,
    calculate_state_duty,
    court_order_exclusion_reason,
    generate_court_orders_zip_result,
    load_base_data,
    load_jurisdiction_data,
    load_static_data,
    validate_base_data,
    validate_jurisdiction_data,
)


def _build_registry(path: Path, rows: list[list[object]]) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Реестр"
    worksheet.append(["Лицевой счет", "ФИО", "Адрес", "Период задолженности", "Сумма долга"])
    for row in rows:
        worksheet.append(row)
    workbook.save(path)


def _build_static_data(path: Path) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.append(["Номер обьекта ", "Адрес", "Компания", "Реквезиты ", "Протокол", "Ставка", "Подсудность"])
    worksheet.append(
        [
            1297,
            "г. Москва, Тестовая улица, д. 1",
            "ТЕСТ ООО",
            (
                'Сокращенное наименование ООО "ТЕСТ"\n'
                "ИНН 7700000000\n"
                "КПП 770001001\n"
                "ОГРН 1200000000000\n"
                "Юридический адрес 127000, г. Москва, Тестовая ул., д. 1 "
                "Генеральный директор Иванов Иван Иванович"
            ),
            "протокол от 01 января 2026 года",
            950,
            "Мировому судье судебного участка № 1",
        ]
    )
    workbook.save(path)


def _build_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("{{ court_name }}")
    document.add_paragraph("Взыскатель: {{ claimant_name }}")
    document.add_paragraph("Должник: {{ debtor_name }}")
    document.add_paragraph("Машино-место № {{ parking_place_number }}")
    document.add_paragraph("Задолженность: {{ debt_amount }}")
    document.add_paragraph("Госпошлина: {{ state_duty }}")
    document.save(path)


class CourtOrdersGenerationTests(unittest.TestCase):
    def test_load_static_data_accepts_current_workbook_headers(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            static_path = Path(temp_dir) / "static.xlsx"
            _build_static_data(static_path)

            data = load_static_data(static_path)

        self.assertIn("1297", data)
        self.assertEqual(data["1297"].claimant_name, 'ООО "ТЕСТ"')
        self.assertEqual(data["1297"].inn, "7700000000")

    def test_calculate_state_duty_for_court_order_uses_half_of_claim_fee(self) -> None:
        self.assertEqual(calculate_state_duty(Decimal("50225.00")), Decimal("2000.00"))
        self.assertEqual(calculate_state_duty(Decimal("150000.00")), Decimal("2750.00"))
        self.assertEqual(calculate_state_duty(Decimal("140001.00")), Decimal("2600"))

    def test_court_debt_uses_only_full_months(self) -> None:
        months, amount = calculate_court_debt(Decimal("12550"), Decimal("1300"))

        self.assertEqual(months, 9)
        self.assertEqual(amount, Decimal("11700"))
        self.assertEqual(
            calculate_court_period("01.01.2026 - 30.06.2026", months),
            "01.10.2025 - 01.06.2026",
        )

    def test_court_order_amount_boundaries_are_inclusive(self) -> None:
        self.assertIsNotNone(court_order_exclusion_reason(4, Decimal("4999")))
        self.assertIsNone(court_order_exclusion_reason(5, Decimal("5000")))
        self.assertIsNone(court_order_exclusion_reason(500, Decimal("500000")))
        self.assertIsNotNone(court_order_exclusion_reason(501, Decimal("500001")))

    def test_generate_court_orders_zip_result_creates_docx_for_known_objects(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            registry_path = temp / "registry.xlsx"
            static_path = temp / "static.xlsx"
            template_path = temp / "template.docx"
            output_zip_path = temp / "court_orders.zip"
            _build_registry(
                registry_path,
                [
                    ["12970001", "Иванов Иван", "Москва, машиноместо № 0001", "01.01.2026 - 31.01.2026", 10000],
                    ["99990001", "Петров Петр", "Москва, машиноместо № 0001", "01.01.2026 - 31.01.2026", 2000],
                ],
            )
            _build_static_data(static_path)
            _build_template(template_path)

            result = generate_court_orders_zip_result(
                registry_path=str(registry_path),
                template_path=str(template_path),
                static_data_path=str(static_path),
                output_zip_path=str(output_zip_path),
                application_date="01.03.2026",
            )

            with ZipFile(output_zip_path) as archive:
                names = archive.namelist()
                docx_name = next(name for name in names if name.endswith(".docx"))
                archive.extract(docx_name, temp)

            generated_doc = Document(str(temp / docx_name))
            generated_text = "\n".join(paragraph.text for paragraph in generated_doc.paragraphs)

        self.assertEqual(result.documents_count, 1)
        self.assertEqual(result.skipped_count, 1)
        self.assertEqual(result.total_debt_amount, Decimal("9500.00"))
        self.assertEqual(len([name for name in names if name.endswith(".docx")]), 1)
        self.assertIn("errors.xlsx", names)
        self.assertIn("Иванов Иван", generated_text)
        self.assertIn("ООО \"ТЕСТ\"", generated_text)
        self.assertIn("9 500,00", generated_text)

    def test_separate_base_and_jurisdiction_files_are_joined_by_object(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            base_path = temp / "base.xlsx"
            jurisdiction_path = temp / "jurisdiction.xlsx"
            _build_static_data(base_path)
            _build_static_data(jurisdiction_path)

            base = load_base_data(base_path)
            jurisdiction = load_jurisdiction_data(jurisdiction_path)

        self.assertEqual(set(base), {"1297"})
        self.assertEqual(set(jurisdiction), {"1297"})
        self.assertEqual(jurisdiction["1297"].court, "Мировому судье судебного участка № 1")

    def test_validation_reports_missing_court_address(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            static_path = Path(temp_dir) / "static.xlsx"
            _build_static_data(static_path)

            base_report = validate_base_data(static_path)
            jurisdiction_report = validate_jurisdiction_data(static_path)

        self.assertEqual(base_report.objects_count, 1)
        self.assertEqual(base_report.warning_counts, {})
        self.assertEqual(jurisdiction_report.warning_counts, {"без адреса суда": 1})

    def test_duplicate_object_code_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            static_path = Path(temp_dir) / "static.xlsx"
            _build_static_data(static_path)
            workbook = load_workbook(static_path)
            worksheet = workbook.active
            worksheet.append(list(worksheet.iter_rows(min_row=2, max_row=2, values_only=True))[0])
            workbook.save(static_path)

            with self.assertRaisesRegex(ValueError, "повторяется код объекта 1297"):
                load_base_data(static_path)

    def test_skipped_rows_are_written_to_errors_workbook(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            registry_path = temp / "registry.xlsx"
            static_path = temp / "static.xlsx"
            template_path = temp / "template.docx"
            output_zip_path = temp / "court_orders.zip"
            _build_registry(
                registry_path,
                [["99990001", "Петров Петр", "Москва", "01.01.2026 - 31.01.2026", 2000]],
            )
            _build_static_data(static_path)
            _build_template(template_path)

            result = generate_court_orders_zip_result(
                registry_path=str(registry_path),
                template_path=str(template_path),
                static_data_path=str(static_path),
                output_zip_path=str(output_zip_path),
            )
            with ZipFile(output_zip_path) as archive:
                names = archive.namelist()

        self.assertEqual(result.documents_count, 0)
        self.assertEqual(result.skipped_count, 1)
        self.assertEqual(names, ["errors.xlsx"])


if __name__ == "__main__":
    unittest.main()
